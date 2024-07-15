import os
from dotenv import load_dotenv
from flask import Flask, request, jsonify
from datetime import datetime
from docx import Document
from docx.shared import Pt
import convertapi
import json
from signature import *

load_dotenv()

document_option1 = os.getenv("DOCUMENT_NAME_OPTION1")
document_option2 = os.getenv("DOCUMENT_NAME_OPTION2")
convertapi_secret_key = os.getenv("CONVERTAPI_SECRET_KEY")

app = Flask(__name__)

counter = 1
prev_date = ""

@app.route('/zap-webhook', methods=['POST'])
def webhook():
    data = request.json
    if data:
        global counter
        global prev_date

        print(f'Received data: {data}')

        with open("webhook.json", "w") as file:
            json.dump(data, file, indent=4)

        if datetime.today().strftime("%m%d%y") == prev_date:
            pass
        else:
            counter = 1
        
        json_data = extract_data(data["body"])
        if json_data == "Not found message":
            return jsonify({"status": "failure", "message": "No data received"}), 400
        else:
            personal_data, prev_date = find_data(json_data[0], data["date"])
            doc_file = replace_data(personal_data, json_data[0]['payment_method'])
            convert_pdf(doc_file)
            boldsign(personal_data, json_data[0]['payment_method'])

            response = {
                "status": "success",
                "message": "Data received successfully!"
            }

            counter += 1

            return jsonify(response), 200
    else:
        return jsonify({"status": "failure", "message": "No data received"}), 400

def extract_data(text):
    final_data = []
    if text:
        splited_texts = text.split("\n\n")
        first_name = splited_texts[0].split(": ")[1]
        last_name = splited_texts[1].split(": ")[1]
        address = splited_texts[2].split(": ")[1]
        city = splited_texts[3].split(": ")[1]
        state = splited_texts[4].split(": ")[1]
        zip = splited_texts[5].split(": ")[1]
        home_phone = splited_texts[6].split(": ")[1]
        cell_phone = splited_texts[7].split(": ")[1]
        email = splited_texts[8].split(": ")[1]
        agreed = splited_texts[9]
        emergency_name = splited_texts[10].split(": ")[1]
        emergency_phone = splited_texts[11].split(": ")[1]
        case_state = splited_texts[12].split(": ")[1]
        case_type = splited_texts[13].split(": ")[1]
        firm_name = splited_texts[14].split(": ")[1]
        attorney_name = splited_texts[15].split(": ")[1]
        attorney_phone = splited_texts[16].split(": ")[1]
        attorney_email = splited_texts[17].split(": ")[1]
        amount = splited_texts[18].split(": ")[1]
        payment_method = splited_texts[19].split(": ")[1]
        description = splited_texts[20].split(": ")[1]

        final_data.append({
            "first_name": first_name,
            "last_name": last_name,
            "address": address,
            "city": city,
            "state": state,
            "zip": zip,
            "home_phone": home_phone,
            "cell_phone": cell_phone,
            "email": email,
            "agreed": agreed,
            "emergency_name": emergency_name,
            "emergency_phone": emergency_phone,
            "case_state": case_state,
            "case_type": case_type,
            "firm_name": firm_name,
            "attorney_name": attorney_name,
            "attorney_phone": attorney_phone,
            "attorney_email": attorney_email,
            "amount": amount,
            "payment_method": payment_method,
            "description": description
        })
        return final_data
    else:
        return "Not found message"

def replace_word_in_paragraphs(paragraphs, old_word, new_word):
    for para in paragraphs:
        if old_word in para.text:
            para.text = para.text.replace(old_word, new_word)
            for run in para.runs:
                run.font.size = Pt(11)

def replace_word_in_tables(tables, old_word, new_word):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_word in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_word, new_word)

def find_data(json_data, date):
    personal_data = []

    client_name = json_data['first_name'] + ' ' + json_data['last_name']
    print(f'<Client_Name> ---> {client_name}')

    client_email = json_data['email']
    print(f'<Client_Email> ---> {client_email}')

    firm_name = json_data['firm_name']
    print(f'<Firm_Name> ---> {firm_name}')
    
    attorney_name = json_data['attorney_name']
    print(f'<Lawyer_Name> ---> {attorney_name}')
    
    firm_address = json_data['attorney_email']
    print(f'<Lawyer_Email> ---> {firm_address}')

    date_obj = datetime.strptime(date, "%a, %d %b %Y %H:%M:%S %z")
    loan_date = date_obj.strftime("%m/%d/%Y")
    print(f'<Loan_Date> ---> {loan_date}')

    today = datetime.today()
    case_id = today.strftime("%m%d%y") + str(f"{counter:02}")
    print(f'<Case_ID> ---> {case_id}')

    loan_amount = json_data['amount']
    print(f'<Loan_Amount> ---> {loan_amount}')

    handling_fee = "0"
    print(f'<Handling_Fee> ---> {handling_fee}')

    client_phone = json_data['home_phone']
    print(f'<Client_Phone> ---> {client_phone}')

    
    q1_interest = int(round(int(json_data['amount']) * 1.15, 2))
    print(f'<Q1_Interest> ---> {q1_interest}')

    q2_interest = int(round(int(json_data['amount']) * 1.3, 2))
    print(f'<Q2_Interest> ---> {q2_interest}')

    q3_interest = int(round(int(json_data['amount']) * 1.45, 2))
    print(f'<Q3_Interest> ---> {q3_interest}')

    q4_interest = int(round(int(json_data['amount']) * 1.6, 2))
    print(f'<Q4_Interest> ---> {q4_interest}')

    client_street = json_data['address']
    print(f'<Client_Street> ---> {client_street}')

    client_citystate = json_data['city'] + ", " + json_data['state'] + " " + json_data['zip']
    print(f'<Client_CityState> ---> {client_citystate}')

    personal_data.append({
        "<Client_Name>" : client_name,
        "<Client_Email>" : client_email,
        "<Firm_Name>" : firm_name,
        "<Lawyer_Name>" : attorney_name,
        "<Lawyer_Email>" : firm_address,
        "<Loan_Date>" : loan_date,
        "<Case_ID>" : case_id,
        "<Loan_Amount>" : loan_amount,
        "<Handling_Fee>" : handling_fee,
        "<Client_Phone>" : client_phone,
        "<Q1_Interest>" : str(q1_interest),
        "<Q2_Interest>" : str(q2_interest),
        "<Q3_Interest>" : str(q3_interest),
        "<Q4_Interest>" : str(q4_interest),
        "<Client_Street>" : client_street,
        "<Client_CityState>" : client_citystate
    })

    return personal_data, today.strftime("%m%d%y")

def replace_data(personal_data, payment_method):
    if "Zelle" in payment_method:
        doc = Document(document_option1)
    else:
        doc = Document(document_option2)

    for key, value in personal_data[0].items():
        replace_word_in_paragraphs(doc.paragraphs, key, value)
        replace_word_in_tables(doc.tables, key, value)

    for paragraph in doc.paragraphs:
        if f"Seller: {personal_data[0]['<Client_Name>']}" in paragraph.text or f"Purchase Price: ${personal_data[0]['<Loan_Amount>']}" in paragraph.text or f"Advance Amount Due to Seller: ${personal_data[0]['<Loan_Amount>']}" in paragraph.text:
            for run in paragraph.runs:
                run.font.size = Pt(12)

    output_file = f"{personal_data[0]["<Client_Name>"]}.docx"
    doc.save(output_file)
    print(f"Words replaced and saved to {personal_data[0]["<Client_Name>"]}.docx")
    
    return output_file

def convert_pdf(file_path):
    pdf_file = file_path.split('.')[0] + '.pdf'
    convertapi.api_secret = convertapi_secret_key
    convertapi.convert('pdf', {
        'File': file_path
    }, from_format = 'doc').save_files(pdf_file)
    os.remove(file_path)
    return pdf_file

if __name__ == '__main__':
    app.run(port=5002, debug=True)